#!/usr/bin/env python3
"""
Utility functions for handling attachments in comments.
Used by both pipeline.py and discover_stances.py
"""

import collections
import os
import re
import base64
import mimetypes
import logging
import requests
from typing import Dict, Any, Tuple, Optional
from dotenv import load_dotenv
import litellm
litellm.drop_params = True  # drop params a model does not support (e.g. temperature on GPT-5 reasoning models)

load_dotenv()

logger = logging.getLogger(__name__)


def is_gibberish(text: str, min_chars: int = 10, max_control: float = 0.02,
                 min_printable: float = 0.80, min_alpha: float = 0.10) -> bool:
    """True when extracted text looks like a failed extraction rather than content.

    Judges character classes, not vocabulary. The previous version required a
    minimum share of common English words, which discarded legitimate documents
    whose text is mostly names, places or tabular data: a real 21-page petition
    with a signatory list scored 4.3% common words (its top terms were state
    abbreviations — ca, tx, pa, ga, fl) and was thrown away even though the
    prose was clean. That check also hard-coded terms from one docket into
    shared code, which the rest of this tool deliberately avoids.

    A genuinely failed extraction does not look like prose at the character
    level. A PDF whose fonts carry no usable cmap yields raw control bytes: one
    real example measured 58% control characters and 5.6% letters, against
    0% and 73.5% for the petition above. That gap is wide, language-agnostic,
    and needs no word list.
    """
    stripped = text.strip() if text else ''
    if len(stripped) < min_chars:
        return True

    n = len(text)
    control = sum(1 for c in text if (ord(c) < 32 and c not in '\t\n\r') or ord(c) == 127)
    if control / n > max_control:
        logger.warning(f"Gibberish: {control / n:.0%} control characters in {text[:60]!r}")
        return True

    printable = sum(1 for c in text if c.isprintable() or c in '\t\n\r')
    if printable / n < min_printable:
        logger.warning(f"Gibberish: only {printable / n:.0%} printable in {text[:60]!r}")
        return True

    alpha = sum(1 for c in text if c.isalpha())
    if alpha / n < min_alpha:
        logger.warning(f"Gibberish: only {alpha / n:.0%} letters in {text[:60]!r}")
        return True

    return False


def extract_text_with_gemini(file_path: str) -> str:
    """Extract text from images using OpenAI vision via LiteLLM.

    NOTE: The function name is legacy (formerly used Google Gemini multimodal).
    The implementation now uses OpenAI's gpt-5.4-mini vision through LiteLLM.
    Only image files are OCR'd here; PDFs and other types are skipped
    (text-based PDFs are handled by the PyMuPDF path in extract_text_from_file).
    """
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        logger.debug("OPENAI_API_KEY not found, skipping vision extraction")
        return ""

    # Check file size (skip large files)
    file_size = os.path.getsize(file_path)
    if file_size > 5 * 1024 * 1024:  # 5MB limit
        logger.warning(f"File too large for vision extraction: {file_path}")
        return ""

    # Determine MIME type; only images are supported by this path
    mime, _ = mimetypes.guess_type(file_path)
    image_mimes = {
        "image/png", "image/jpeg", "image/gif", "image/webp",
    }
    if mime not in image_mimes:
        logger.debug(
            f"Skipping vision extraction for non-image file {file_path} "
            f"(mime={mime}); PDFs/text handled by other extraction paths"
        )
        return ""

    try:
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")

        resp = litellm.completion(
            model="gpt-5.4-mini",
            messages=[{"role": "user", "content": [
                {"type": "text", "text": "Extract all text from this document. Return only the raw text content. If there is no readable text, return exactly the word EMPTY and nothing else."},
                {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
            ]}],
            temperature=0.0,
        )
        text = resp.choices[0].message.content or ""
        text = text.strip()

        # Check for the EMPTY sentinel we asked for (or empty response)
        if not text or text.upper() == 'EMPTY':
            logger.info(f"Vision extraction found no text in {file_path}")
            return ""

        if is_gibberish(text):
            logger.warning(f"Vision extraction returned gibberish for {file_path}, discarding")
            return ""
        return text

    except Exception as e:
        logger.warning(f"Vision extraction failed for {file_path}: {e}")
        return ""

# Order to try the files of one logical attachment. regulations.gov stores the
# submitter's original alongside a PDF rendition of the same document, so these
# are alternative encodings of one thing, not separate content: .docx + .pdf
# alone covers 688 of 862 multi-file comments on OMB-2026-0034.
#
# .docx leads because it cannot be a scan, cannot carry a broken font cmap, and
# needs no OCR — the three ways PDFs failed here. Measured against 20 real
# pairs it recovers 97-99% of the PDF's text, the shortfall being page numbers
# and running headers the renderer adds. Images come last: they cost a vision
# call, so they are only reached when nothing else exists.
_FORMAT_PREFERENCE = ['.docx', '.doc', '.pdf', '.txt', '.rtf', '.html', '.htm',
                      '.png', '.jpg', '.jpeg', '.gif', '.webp']


def _attachment_stem(filename: str) -> str:
    """Filename with our attachment_N_ prefix and extension removed.

    Siblings of one upload share this stem — true for 846 of 862 multi-file
    comments — so it is what groups alternative encodings together.
    """
    base = re.sub(r'\.[A-Za-z0-9]+$', '', filename)
    return re.sub(r'^attachment_\d+_', '', base).strip().lower()


def _rank(filename: str) -> int:
    ext = os.path.splitext(filename)[1].lower()
    return _FORMAT_PREFERENCE.index(ext) if ext in _FORMAT_PREFERENCE else len(_FORMAT_PREFERENCE)


def group_attachments(entries: list) -> list:
    """Group (url, filename) pairs by stem; order each group by preference.

    Returns a list of groups. Extracting the first entry of a group that yields
    text is enough — the rest are the same document in another format.
    """
    groups = collections.OrderedDict()
    for url, filename in entries:
        groups.setdefault(_attachment_stem(filename), []).append((url, filename))
    return [sorted(v, key=lambda uf: _rank(uf[1])) for v in groups.values()]


def _pages_needing_ocr(doc, min_chars: int = 100) -> list:
    """Page indexes with visible content but no usable text, so only OCR can read them.

    Three real cases reach this, and an image test alone catches only the first:
      - an image-only scan (no text layer at all)
      - text drawn as vector outlines, which carries no text and no image either
      - a font with no usable cmap, which yields control bytes rather than nothing

    So the test is "has ink but no readable text", not "has a big image". Pages
    that are genuinely blank have no ink and are skipped, since OCR would only
    spend a call to confirm they are empty.
    """
    out = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if len(text) >= min_chars and not is_gibberish(text):
            continue
        if page.get_image_info() or page.get_drawings():
            out.append(i)
    return out


def ocr_scanned_pdf(file_path: str, max_pages: int = 20) -> str:
    """OCR the scan-like pages of a PDF by rasterising them and using the vision path.

    Capped at max_pages so one long scan can't run away with the budget.
    """
    import tempfile

    import fitz

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.error(f"Cannot open PDF for OCR {file_path}: {e}")
        return ""

    with doc:
        targets = _pages_needing_ocr(doc)
        if not targets:
            return ""
        if len(targets) > max_pages:
            logger.warning(
                f"{os.path.basename(file_path)}: {len(targets)} scanned pages, "
                f"OCRing the first {max_pages}")
            targets = targets[:max_pages]

        logger.info(f"  {os.path.basename(file_path)}: OCRing {len(targets)} scanned page(s)")
        chunks = []
        for i in targets:
            tmp = None
            try:
                # 200 dpi is enough for body text and keeps the PNG under the
                # vision path's 5 MB ceiling.
                pix = doc[i].get_pixmap(dpi=200)
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as fh:
                    tmp = fh.name
                pix.save(tmp)
                page_text = extract_text_with_gemini(tmp)
                if page_text:
                    chunks.append(page_text)
            except Exception as e:
                logger.warning(f"  OCR failed on page {i + 1} of {file_path}: {e}")
            finally:
                if tmp and os.path.exists(tmp):
                    os.remove(tmp)

    return '\n'.join(chunks)


def download_attachment(attachment_url: str, output_path: str) -> bool:
    """Download an attachment file."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'}
        response = requests.get(attachment_url, stream=True, timeout=30, headers=headers)
        response.raise_for_status()
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        return True
    except Exception as e:
        logger.error(f"Failed to download {attachment_url}: {e}")
        return False

def extract_text_from_file(file_path: str, use_gemini: bool = False) -> str:
    """Extract text from various file types."""
    import fitz  # PyMuPDF
    import docx

    # Try Gemini first if enabled and available
    if use_gemini:
        gemini_text = extract_text_with_gemini(file_path)
        if gemini_text:
            return gemini_text

    text = ""
    if file_path.lower().endswith('.pdf'):
        # PyMuPDF preserves visual reading order (position-aware blocks), unlike
        # PyPDF2 which follows raw content-stream order and garbles multi-column
        # layouts (e.g. two-column signature blocks come out one word per line).
        try:
            with fitz.open(file_path) as doc:
                text = '\n'.join(page.get_text() for page in doc)
                scanned = _pages_needing_ocr(doc)
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            return ""

        # A broken font cmap yields control bytes rather than nothing. Drop that
        # before merging so it can't contaminate the OCR result we are about to
        # fetch, or trip the gibberish check on the combined string.
        if text.strip() and is_gibberish(text):
            logger.info(
                f"{os.path.basename(file_path)}: text layer is unreadable, "
                f"using OCR instead")
            text = ""

        # Pages with ink but no readable text need OCR, or they extract to nothing.
        if scanned:
            ocr_text = ocr_scanned_pdf(file_path)
            if ocr_text:
                text = f"{text}\n{ocr_text}".strip() if text.strip() else ocr_text
            else:
                logger.warning(
                    f"{os.path.basename(file_path)}: {len(scanned)} unreadable page(s) "
                    f"yielded no OCR text")

    elif file_path.lower().endswith(('.doc', '.docx')):
        try:
            doc = docx.Document(file_path)
            parts = [p.text for p in doc.paragraphs]
            # Tables are not paragraphs. Reading only doc.paragraphs dropped
            # 3,336 characters from one real 53k-character submission, because
            # its signatory block was a table.
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(' '.join(cells))
            text = '\n'.join(parts)
        except Exception as e:
            logger.error(f"Failed to extract text from DOC {file_path}: {e}")
            return ""
    
    elif file_path.lower().endswith('.txt'):
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except Exception as e:
            logger.error(f"Failed to read text file {file_path}: {e}")
            return ""

    elif file_path.lower().endswith(('.html', '.htm')):
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
        except Exception as e:
            logger.error(f"Failed to extract text from HTML {file_path}: {e}")
            return ""

    else:
        logger.warning(f"Unsupported file type: {file_path}")
        return ""

    # Check for gibberish before returning
    if is_gibberish(text):
        logger.warning(f"Gibberish detected in local extraction of {file_path}, discarding")
        return ""
    return text


def reextract_attachment_text(comment_id: str, attachments_dir: str = 'attachments') -> Optional[str]:
    """Re-run extraction on a comment's already-downloaded PDF attachment(s),
    refresh the on-disk `.extracted.txt` cache, and return the combined text.

    Used to pick up extractor improvements (e.g. the PyPDF2 -> PyMuPDF swap) for
    specific comments without re-downloading or re-processing the whole corpus.
    Returns None if no cached PDF attachment is found for this comment.
    """
    comment_dir = os.path.join(attachments_dir, comment_id)
    if not os.path.isdir(comment_dir):
        return None

    pdf_paths = sorted(
        os.path.join(comment_dir, f) for f in os.listdir(comment_dir) if f.lower().endswith('.pdf')
    )
    if not pdf_paths:
        return None

    parts = []
    for pdf_path in pdf_paths:
        text = extract_text_from_file(pdf_path)
        text_cache_path = f"{pdf_path}.extracted.txt"
        try:
            with open(text_cache_path, 'w', encoding='utf-8') as f:
                f.write(text or "")
        except Exception as e:
            logger.warning(f"Failed to refresh text cache {text_cache_path}: {e}")
        if text and text.strip():
            parts.append(text.strip())

    return '\n\n'.join(parts)


def process_attachments(comment_data: Dict[str, Any], attachments_dir: str,
                       attachment_col: str = 'Attachment Files',
                       download_missing: bool = True,
                       use_gemini: bool = False) -> Tuple[str, Dict[str, Any]]:
    """
    Download and process attachments for a comment, return combined text and processing status.
    
    Args:
        comment_data: Dictionary containing comment data
        attachments_dir: Base directory for storing attachments
        attachment_col: Name of the column containing attachment URLs
        download_missing: Whether to download attachments that don't exist locally
        use_gemini: Whether to use Gemini API for text extraction (requires GEMINI_API_KEY)
    
    Returns:
        Tuple of (combined_text, processing_status)
    """
    comment_id = comment_data.get('Document ID', 'Unknown')
    logger.info(f"=== PROCESSING ATTACHMENTS FOR {comment_id} ===")
    
    if attachment_col not in comment_data or not comment_data[attachment_col]:
        logger.info(f"  No attachments found for {comment_id}")
        return "", {"total": 0, "processed": 0, "failed": 0, "failures": []}
    
    attachment_urls = comment_data[attachment_col].split(',')
    logger.info(f"  Found {len(attachment_urls)} attachment URLs")
    combined_attachment_text = []
    processing_status = {
        "total": len([url for url in attachment_urls if url.strip()]),
        "processed": 0,
        "failed": 0,
        "failures": []
    }
    
    # Create directory for this comment's attachments
    comment_id = (comment_data.get('Document ID') or 
                 comment_data.get('Comment ID') or 
                 'unknown_comment')
    comment_attachment_dir = os.path.join(attachments_dir, comment_id)
    
    # One upload is stored as several files (the submitter's original plus a PDF
    # rendition). Group them and read only the first that yields text, instead of
    # extracting every copy of the same document.
    entries = []
    for i, url in enumerate(attachment_urls):
        url = url.strip()
        if not url:
            continue
        filename = f"attachment_{i+1}_{url.split('/')[-1]}"
        if '.' not in filename:
            filename += '.pdf'  # Default extension
        entries.append((url, filename))

    groups = group_attachments(entries)
    processing_status["groups"] = len(groups)
    processing_status["skipped_duplicate"] = 0
    had_pdf = any(f.lower().endswith('.pdf') for _, f in entries)

    for group in groups:
        got_text = False
        for url, filename in group:
            if got_text:
                # An earlier format in this group already gave us the document.
                logger.info(f"  Skipping {filename}: same document as a file already read")
                processing_status["skipped_duplicate"] += 1
                continue

            file_path = os.path.join(comment_attachment_dir, filename)
            text_cache_path = os.path.join(comment_attachment_dir, f"{filename}.extracted.txt")

            # An EMPTY cache file is not treated as an answer: it used to
            # permanently pin scanned PDFs at no text, so a fix to the extractor
            # could never reach them. Falling through re-extracts, and a
            # successful extraction then caches real text.
            if os.path.exists(text_cache_path):
                try:
                    with open(text_cache_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    if text.strip():
                        combined_attachment_text.append(text.strip())
                        logger.info(f"  SUCCESS: Loaded {len(text)} characters from cache "
                                    f"({filename})")
                        processing_status["processed"] += 1
                        got_text = True
                        continue
                    logger.info(f"  Cached text for {filename} is empty, re-extracting")
                except Exception as e:
                    logger.warning(f"  Failed to load cached text: {e}")

            # Check if attachment file already exists
            if os.path.exists(file_path):
                logger.info(f"  Attachment {filename} already exists, skipping download")
            else:
                if not download_missing:
                    logger.info(f"  Skipping download of {filename} (download_missing=False)")
                    processing_status["failed"] += 1
                    processing_status["failures"].append(
                        {"filename": filename, "reason": "skipped_download"})
                    continue

                logger.info(f"  Downloading attachment: {filename}")
                if not download_attachment(url, file_path):
                    processing_status["failed"] += 1
                    processing_status["failures"].append(
                        {"filename": filename, "reason": "download_failed"})
                    continue

            # An image is only reached when no richer format exists in the group,
            # so read it rather than skipping: that is the one case where the
            # picture IS the comment. Cheap, because it is rare.
            is_image = os.path.splitext(filename)[1].lower() in {
                '.png', '.jpg', '.jpeg', '.gif', '.webp'}

            logger.info(f"  Extracting text from {filename}...")
            text = extract_text_from_file(file_path, use_gemini=use_gemini or is_image)

            os.makedirs(os.path.dirname(text_cache_path), exist_ok=True)
            try:
                with open(text_cache_path, 'w', encoding='utf-8') as f:
                    f.write(text or "")
            except Exception as e:
                logger.warning(f"  Failed to save text cache: {e}")

            if not text or not text.strip():
                logger.warning(f"  No text extracted from {filename}")
                processing_status["failed"] += 1
                processing_status["failures"].append(
                    {"filename": filename, "reason": "no_text_extracted"})
                continue

            combined_attachment_text.append(text.strip())
            logger.info(f"  SUCCESS: Extracted {len(text)} characters from {filename}")
            processing_status["processed"] += 1
            got_text = True

        if not got_text:
            # Every format of this upload failed. Say so loudly, and say whether a
            # PDF was even on offer — regulations.gov renders one for all but
            # ~1 in 2,600 uploads, so "no PDF" means this docket breaks the
            # pattern the extractor is built around.
            names = ', '.join(f for _, f in group)
            logger.warning(
                f"  UNREADABLE ATTACHMENT on {comment_id}: no text from any of [{names}]"
                + ("" if had_pdf else "  (no PDF rendition offered — unusual, check this docket)"))
            processing_status["failures"].append(
                {"filename": names, "reason": "no_pdf_rendition" if not had_pdf else "all_formats_empty"})

    logger.info(f"=== ATTACHMENT PROCESSING COMPLETE FOR {comment_id} ===")
    logger.info(f"  Status: {processing_status}")
    logger.info(f"  Total text extracted: {len(''.join(combined_attachment_text))} characters")
    
    return "\n\n--- ATTACHMENT ---\n\n".join(combined_attachment_text), processing_status